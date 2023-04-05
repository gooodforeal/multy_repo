from docx import Document
import datetime


def main():
    with open("data.txt", "r", encoding="UTF-8") as f:
        lines = f.readlines()
    print(lines)
    template = 'template.docx'
    result = 'result.docx'
    data = {
    "{{EMPLOEE_POSITION}}": lines[0],
    "{{EMPLOEE_FULLNAME}}": lines[1],
    "{{EMPLOEE_SHORTNAME}}": lines[2],
    "{{START_DATE}}": lines[3],
    "{{FINISH_DATE}}": lines[4],
    "{{DURATION}}": lines[5],
    "{{DATE}}": str(datetime.datetime.today().date()).replace("-", "."),
    }
    template_doc = Document(template)
    for key, value in data.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, key, value)
            template_doc.save(result)
        for row in template_doc.tables[0].rows:
            for cell in row.cells:
                replace_text(cell, key, value)
                template_doc.save(result)


def replace_text(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


main()
